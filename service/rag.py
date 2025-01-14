import os
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.schema import Document as LangchainDocument
from docx import Document
from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain.memory import ConversationSummaryMemory
from langchain.prompts import PromptTemplate
from langchain.retrievers.document_compressors import LLMChainFilter
from langchain.retrievers import ContextualCompressionRetriever
from langsmith import traceable
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain.prompts import ChatPromptTemplate, HumanMessagePromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from service.custom_retriever import DentistPatientRetriever


def load_docx_from_dir(document_direction):
    """
    Load .docx files from a directory and extract their content into a dictionary.

    Args:
        document_direction (str): Directory path containing .docx files.

    Returns:
        Dict[str, List[str]]: Dictionary with file names as keys and lists of extracted text as values.
    """
    document_texts = {}

    for file_name in os.listdir(document_direction):
        if file_name.endswith('.docx'):
            file_path = os.path.join(document_direction, file_name)
            doc = Document(file_path)

            # Parse the document to extract relevant text
            parsed_text = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:  # Include all lines, not just those with ":"
                    parsed_text.extend([line.strip() for line in text.splitlines() if line.strip()])

            document_texts[file_name.split('.')[0]] = parsed_text
    print(document_texts)
    return document_texts



def split_conversations_into_chunks(document_texts, group_size = 6):
    """
    Split document text into chunks of a specified size.

    Args:
        document_texts (Dict[str, List[str]]): Dictionary with file names as keys and lists of extracted text as values.
        group_size (int): Number of items per chunk.

    Returns:
        Dict[str, List[List[str]]]: Dictionary with file names as keys and chunked lists of text as values.
    """
    chunked_texts = {}

    for doc_name, text_list in document_texts.items():
        chunks = [text_list[i:i + group_size] for i in range(0, len(text_list), group_size)]
        chunked_texts[doc_name] = chunks

    return chunked_texts

def store_data(OPENAI_API_KEY):
    
    document_direction = "./chatstreamlit/src/document"
    document_texts = load_docx_from_dir(document_direction)

    chunked_data=split_conversations_into_chunks(document_texts, group_size = 6)

    split_documents = []
    for doc_name, chunks in chunked_data.items():
        for chunk in chunks:
            content = "\n".join(chunk)
            split_documents.append(LangchainDocument(page_content=content))
    print(split_documents)
    # Create embeddings for the documents
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY, model="text-embedding-ada-002")
    db = Chroma.from_documents(documents=split_documents, embedding=embeddings,
                               persist_directory='./chatstreamlit/src/chroma')
    retriever = db.as_retriever()

    print(f'Finished storing data. Retriever: {retriever}')
    return retriever

# def rertrive_data(query, OPENAI_API_KEY, persist_directory):
#     embedding_function = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
#     db = Chroma(persist_directory=persist_directory, embedding_function=embedding_function)
#     # Set up the retriever using max marginal relevance (MMR) search
#     # This limits the number of documents returned to 3 (k=3)
#     retriever = db.as_retriever(search_type="mmr", search_kwargs={"k": 3})
#     llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
#     # Create an LLM-based filter to refine retrieved documents
#     llm_filter = LLMChainFilter.from_llm(llm)
#     # Combine the base retriever and LLM filter into a Contextual Compression Retriever
#     # This enhances the relevance of documents by filtering irrelevant content
#     compression_retriever = ContextualCompressionRetriever(
#         base_compressor=llm_filter, base_retriever=retriever
#     )
#
#     qa_chain = RetrievalQA.from_chain_type(
#         llm=llm,
#         retriever=compression_retriever,
#         return_source_documents=True
#     )
#     response = qa_chain(query)
#     print(f'The response is {response}')
#
#     return response['result']

@traceable
def rertrive_data(query, OPENAI_API_KEY, persist_directory):
    embedding_function = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    # Initialize custom retriever
    custom_retriever = DentistPatientRetriever(
        embedding_function=embedding_function,
        persist_directory=persist_directory
    )
    
    
    llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
    # Create an LLM-based filter to refine retrieved documents
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=custom_retriever,
        return_source_documents=True
    )
    response = qa_chain(query)
    print(f'The response is {response}')
    
    return response['result']

def retrive_memory_and_prompt(question,OPENAI_API_KEY):
    llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
    memory = ConversationSummaryMemory(
        llm=llm,
        memory_key="chat_history",
        return_messages=True,
        max_token_limit=3600,
        output_key="result"
    )
    print(memory)
    embedding_function = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedding_function)

    template = """You are the patient who is going to see a dentist. what you response should base on your personality.
        The conversation will base on scenario. If the message has been told (known_message) to the dentist,
        and when the dentist repeat again the question, you should explain in detail for your previous response.
        When the dentist asks if you have any questions, you might want to inquire about the treatment plan,
        including the duration, specific procedures, and pain management options. Additionally, ask about post-procedure care,
        such as recovery time and any dietary or activity restrictions. Finally, discuss preventive measures,
        follow-up arrangements, and cost and insurance coverage.
        Only respond to the dentist's question without including any unrelated content (in several sentences).
        The entire conversation should revolve around inquiring about detailed patient information before performing any actual dental diagnostic procedures.
        The generated dialogue should be coherent and natural, with seamless transitions.
        As the patient visiting a dentist, follow the scenario below to answer the dentist's question in a few sentences from the patient's perspective.
        It should generate only several sentences and wait for the dentist to respond.
        Remember to keep your response relevant to the dentist's question from the patient's perspective.
    {context}
    Question: {question}
    patient's Answer:"""
    QA_CHAIN_PROMPT = PromptTemplate(input_variables=["context", "question"], template=template, )

    
    
    qa_chain = RetrievalQA.from_chain_type(llm,
                                           retriever=vectordb.as_retriever(),
                                           return_source_documents=True,
                                           memory=memory,
                                           output_key="result",
                                           chain_type_kwargs={"prompt": QA_CHAIN_PROMPT})

    result = qa_chain({"query": question})
    response=result["result"]
    print(f'the response is {response}')

    memory_contents = memory.load_memory_variables({})
    print("\nCurrent Memory Contents:")
    print(memory_contents["chat_history"])
    
    return response

# format the docs
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

def Rag_chain(question, prompt, llm, OPENAI_API_KEY, patient_information, known_message, emotion, scenario):
    embedding_function = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    # vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedding_function)
    # retriever = vectordb.as_retriever()
    custom_retriever = DentistPatientRetriever(
        embedding_function=embedding_function,
        persist_directory=persist_directory
    )
    retrieved_docs = custom_retriever._get_relevant_documents(question)
    context = format_docs(retrieved_docs)
    print(f"\nFormatted Context: {context}")
    rag_chain = (
            {
                "context": RunnablePassthrough(),
                "patient_information": RunnablePassthrough(),
                "known_message": RunnablePassthrough(),
                "dentist_question": RunnablePassthrough(),
                "emotion": RunnablePassthrough(),
                "scenario": RunnablePassthrough()
            }
            | prompt
            | llm
            | StrOutputParser()
    )
    answer = rag_chain.invoke({
        "context": context,
        "patient_information": patient_information,
        "known_message": known_message,
        "dentist_question": question,
        "emotion": emotion,
        "scenario": scenario
    })

    
    print(f'The answer is {answer}')
    return answer

if __name__ == '__main__':
    persist_directory='./chatstreamlit/src/chroma'
    load_dotenv()
    OPENAI_API_KEY=os.getenv("OPENAI_API_KEY")
    
    retriever=store_data(OPENAI_API_KEY)
    # question='Is there anything you not clear and want to ask me today?'
    #
    # llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=OPENAI_API_KEY)
    # prompt="""You are the patient who is going to see a dentist. what you response should base on your personality.
    # The conversation will base on scenario. If the message has been told (known_message) to the dentist,
    # and when the dentist repeat again the question, you should explain in detail for your previous response.
    # When the dentist asks if you have any questions, you might want to inquire about the treatment plan,
    # including the duration, specific procedures, and pain management options. Additionally, ask about post-procedure care,
    # such as recovery time and any dietary or activity restrictions. Finally, discuss preventive measures,
    # follow-up arrangements, and cost and insurance coverage.
    # Only respond to the dentist's question without including any unrelated content (in several sentences).
    # The entire conversation should revolve around inquiring about detailed patient information before performing any actual dental diagnostic procedures.
    # The generated dialogue should be coherent and natural, with seamless transitions.
    # As the patient visiting a dentist, follow the scenario below to answer the dentist's question in a few sentences from the patient's perspective.
    # It should generate only several sentences and wait for the dentist to respond.
    # The answer should remove "Patient:"
    #
    # The information of you is:
    # ###
    # {patient_information}
    # ###
    # The message that you have already told dentist is:
    # ###
    # {known_message}
    # ###
    # The dentist's question is:
    # ###
    # {dentist_question}
    # ###
    # The personality of you is:
    # ###
    # {emotion}
    # ###
    # The scenario of you is:
    # ###
    # {scenario}
    # ###
    #
    # Remember to keep your response relevant to the dentist's question from the patient's perspective.
    # """
    #
    # patient_information = "You are 30 years old with no known allergies."
    # known_message = "You told the dentist that you have been experiencing tooth pain for two weeks."
    # emotion = "You are feeling anxious about the procedure."
    # scenario = "You are at a dentist's office for a root canal consultation."
    #
    # prompt_template = ChatPromptTemplate.from_messages([
    #     HumanMessagePromptTemplate.from_template(prompt)
    # ])
    # response = Rag_chain(
    #     question=question,
    #     prompt=prompt_template,
    #     llm=llm,
    #     OPENAI_API_KEY=OPENAI_API_KEY,
    #     patient_information=patient_information,
    #     known_message=known_message,
    #     emotion=emotion,
    #     scenario=scenario
    # )
